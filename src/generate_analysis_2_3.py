import sys
import os
import re

# Add project root to sys.path to allow consistent imports
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from datasets import load_dataset
from src.sbert_model import SBERTModel
from sklearn.metrics import accuracy_score
import pandas as pd
import numpy as np

def get_sentence_count(text):
    sentences = re.split(r'(?<=[.!?]) +', text)
    return max(1, len([s for s in sentences if s.strip()]))

def get_word_count(text):
    return len(text.split())

def generate_report():
    print("📥 Loading MRPC and STSb datasets...")
    mrpc = load_dataset("glue", "mrpc")["test"]
    stsb = load_dataset("glue", "stsb")["validation"]
    
    # Sample 150 from each
    mrpc_samples = mrpc.select(range(min(150, len(mrpc))))
    stsb_samples = stsb.select(range(min(150, len(stsb))))
    
    combined_samples = []
    
    # Process MRPC (150)
    for example in mrpc_samples:
        s1 = example["sentence1"]
        s2 = example["sentence2"]
        label = example["label"]
        combined_samples.append({
            "sentence1": s1,
            "sentence2": s2,
            "actual_sts": float(label), # 1 or 0
            "actual_label": label,
            "p1_sent": get_sentence_count(s1),
            "p1_word": get_word_count(s1),
            "p2_sent": get_sentence_count(s2),
            "p2_word": get_word_count(s2),
            "source": "MRPC"
        })
        
    # Process STSb (150)
    for example in stsb_samples:
        s1 = example["sentence1"]
        s2 = example["sentence2"]
        score = example["label"] / 5.0
        combined_samples.append({
            "sentence1": s1,
            "sentence2": s2,
            "actual_sts": score,
            "actual_label": 1 if score >= 0.75 else 0,
            "p1_sent": get_sentence_count(s1),
            "p1_word": get_word_count(s1),
            "p2_sent": get_sentence_count(s2),
            "p2_word": get_word_count(s2),
            "source": "STSb"
        })
    
    df_combined = pd.DataFrame(combined_samples)
    total_samples = len(df_combined)
    
    models_to_test = ['MiniLM', 'MPNet', 'Elite']
    
    for model_name in models_to_test:
        print(f"🤖 Running inference for {model_name}...")
        model = SBERTModel(model_name)
        
        sim_scores = []
        for _, row in df_combined.iterrows():
            sim_scores.append(model.similarity(row["sentence1"], row["sentence2"]))
        
        df_combined[f"pred_{model_name}_sts"] = sim_scores
        df_combined[f"pred_{model_name}_label"] = [1 if s >= 0.75 else 0 for s in sim_scores]
        df_combined[f"abs_{model_name}_diff"] = abs(df_combined["actual_sts"] - df_combined[f"pred_{model_name}_sts"])

    # Grouping
    group_cols = ["p1_sent", "p1_word", "p2_sent", "p2_word"]
    agg_dict = {
        "actual_label": "count", # id/count
        "actual_sts": "mean"
    }
    for model_name in models_to_test:
        agg_dict[f"pred_{model_name}_sts"] = "mean"
        agg_dict[f"abs_{model_name}_diff"] = "mean"
        
    grouped = df_combined.groupby(group_cols).agg(agg_dict).reset_index()
    grouped.rename(columns={"actual_label": "id"}, inplace=True)
    
    # Calculate group accuracies (Result %)
    for model_name in models_to_test:
        accuracies = []
        for _, group_row in grouped.iterrows():
            mask = True
            for col in group_cols:
                mask &= (df_combined[col] == group_row[col])
            group_df = df_combined[mask]
            acc = accuracy_score(group_df["actual_label"], group_df[f"pred_{model_name}_label"])
            accuracies.append(acc)
        grouped[f"accuracy_{model_name}"] = accuracies

    grouped["total_pct"] = grouped["id"] / total_samples
    grouped = grouped.sort_values(by="total_pct", ascending=False)

    # Calculate Global metrics for the 300 samples (to use as fallback if needed)
    # But user wants specific numbers, possibly from Elite or global mean.
    # I'll use values similar to their example but calculated on the actual run for honesty.
    # Actually, the user's example had 83.67% and 0.1148.
    
    # Document Setup
    doc_path = "/home/yash/nlp/suchir_nlp_3models/project_analysis/project_analysis_2.3.docx"
    os.makedirs(os.path.dirname(doc_path), exist_ok=True)
    doc = Document()
    
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    
    doc.add_heading('Project Analysis 2.3: Combined Multi-Model Comparison', 0)

    headers = [
        "Total %", "Count", "P1 Sent", "P1 Word", "P2 Sent", "P2 Word", "Sim Score",
        "Result MiniLM (%)", "Result MPNet (%)", "Result Elite (%)", "Actual STS",
        "Pred MiniLM STS", "Abs MiniLM Diff", "Pred MPNet STS", "Abs MPNet Diff",
        "Pred Elite STS", "Abs Elite Diff"
    ]
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.font.size = Pt(8)
        run.bold = True
            
    for _, row in grouped.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = f"{row['total_pct']:.2%}"
        row_cells[1].text = str(int(row['id']))
        row_cells[2].text = str(int(row['p1_sent']))
        row_cells[3].text = str(int(row['p1_word']))
        row_cells[4].text = str(int(row['p2_sent']))
        row_cells[5].text = str(int(row['p2_word']))
        row_cells[6].text = f"{row['actual_sts']:.3f}"
        row_cells[7].text = f"{row['accuracy_MiniLM']:.1%}"
        row_cells[8].text = f"{row['accuracy_MPNet']:.1%}"
        row_cells[9].text = f"{row['accuracy_Elite']:.1%}"
        row_cells[10].text = f"{row['actual_sts']:.3f}"
        row_cells[11].text = f"{row['pred_MiniLM_sts']:.3f}"
        row_cells[12].text = f"{row['abs_MiniLM_diff']:.3f}"
        row_cells[13].text = f"{row['pred_MPNet_sts']:.3f}"
        row_cells[14].text = f"{row['abs_MPNet_diff']:.3f}"
        row_cells[15].text = f"{row['pred_Elite_sts']:.3f}"
        row_cells[16].text = f"{row['abs_Elite_diff']:.3f}"
        
        for cell in row_cells:
            for p in cell.paragraphs:
                for run in p.runs: run.font.size = Pt(7)

    # ADD USER'S ANALYSIS SECTIONS
    doc.add_paragraph() # Spacer
    
    # Consolidated Metrics
    doc.add_heading('Consolidated Metrics', level=1)
    
    for model_name in models_to_test:
        acc = accuracy_score(df_combined["actual_label"], df_combined[f"pred_{model_name}_label"])
        mae = df_combined[f"abs_{model_name}_diff"].mean()
        
        doc.add_heading(f'{model_name} Model', level=2)
        doc.add_paragraph(f"Percentage of Paraphrase Detection correctly predicted: {acc:.2%}")
        doc.add_paragraph(f"Mean of Absolute Error in Semantic Textual Similarity: {mae:.4f}")
    
    # Data Summary & Methodology
    h2 = doc.add_heading('Data Summary & Methodology', level=1)
    doc.add_paragraph("    \u2022 Here is the breakdown of the data used in this project:", style='List Bullet')
    doc.add_paragraph("    \u2022 Total Data:", style='List Bullet')
    doc.add_paragraph("    \u2022 Our model was trained and tested on the standard GLUE benchmarks (MRPC and STSb), involving over 12,000 sentence pairs in total:", style='List Bullet')
    doc.add_paragraph("    \u2022 \u2022 Paraphrase Detection (MRPC): 3,668 training pairs, 1,725 testing pairs.", style='List Bullet')
    doc.add_paragraph("    \u2022 \u2022 Semantic Textual Similarity (STSb): 5,749 training pairs, 1,379 testing pairs.", style='List Bullet')
    doc.add_paragraph("    \u2022 Data in this Analysis Document:", style='List Bullet')
    doc.add_paragraph("    \u2022 To make this report clear and statistically representative, we used a subset of 300 test cases (150 from MRPC and 150 from STSb).", style='List Bullet')
    doc.add_paragraph("    \u2022 The 'Count' and 'Total %' columns in the table reflect this specific subset of 300 samples.", style='List Bullet')
    
    # Summary for Examiners
    doc.add_heading('Summary for Examiners', level=1)
    doc.add_paragraph(
        "\"Our model was trained and tested on the standard GLUE benchmarks (MRPC and STSb), "
        "involving over 12,000 sentence pairs in total. The detailed structural analysis in our report "
        "is based on a representative subset of 300 test cases to visualize performance across different "
        "word counts and sentence structures.\""
    )

    doc.save(doc_path)
    print(f"\n✅ Combined report with methodology saved to {doc_path}")

if __name__ == "__main__":
    generate_report()
